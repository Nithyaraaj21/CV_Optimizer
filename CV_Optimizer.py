import docx
import nltk
import spacy
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.corpus import wordnet

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

# Load pre-trained spaCy model with word vectors
nlp = spacy.load("en_core_web_md")

def read_docx(file_path):
    doc = docx.Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    return content

def preprocess_text(text):
    # Tokenization
    tokens = word_tokenize(text)
    # Lowercasing
    tokens = [word.lower() for word in tokens]
    # Remove punctuation and stopwords
    stop_words = set(stopwords.words('english'))
    tokens = [word for word in tokens if word.isalnum() and word not in stop_words]
    return ' '.join(tokens)

def get_synonyms(word):
    synonyms = set()
    for synset in wordnet.synsets(word):
        for lemma in synset.lemmas():
            synonyms.add(lemma.name())
    return synonyms

def generate_report(matched_points, matching_keywords, similarity_scores, synonyms):
    report_doc = docx.Document()
    report_doc.add_heading('Matched Bullet Points Report', level=1)

    # Add keywords and synonyms to the report
    report_doc.add_heading('Keywords', level=2)
    for keyword in matching_keywords:
        report_doc.add_paragraph(f'- {keyword}')

    report_doc.add_heading('Synonyms', level=2)
    for keyword, synonym_list in synonyms.items():
        synonym_str = ', '.join(synonym_list)
        report_doc.add_paragraph(f'{keyword}: {synonym_str}')

    # Add matched bullet points and their scores to the report
    report_doc.add_heading('Matched Bullet Points', level=2)
    for point, score in zip(matched_points, similarity_scores):
        report_doc.add_paragraph(f'- {point} (Score: {score:.2f})')

    return report_doc

if __name__ == "__main__":
    # File paths for the CV and job description
    cv_path = 'CV_KP.docx'
    job_description_path = 'job_description.docx'
    output_file_path = 'matched_bullet_points.docx'
    report_file_path = 'report.docx'

    # Read the content of the CV and job description
    cv_content = read_docx(cv_path)
    job_description_content = read_docx(job_description_path)

    # Preprocess text
    cv_text = ' '.join(cv_content)
    job_description_text = preprocess_text(' '.join(job_description_content))

    # Calculate Word Movers' Distance similarity using spaCy
    cv_doc = nlp(cv_text)
    job_description_doc = nlp(job_description_text)
    similarity_score = cv_doc.similarity(job_description_doc)

    print("Similarity Score:", similarity_score)

    # Extract keywords from job description
    job_description_keywords = set([token.text.lower() for token in job_description_doc if not token.is_stop and token.is_alpha])

    # Create a new document for matched bullet points
    matched_points = []
    similarity_scores = []
    synonyms = {}

    # Extract bullet points from CV and check for matching keywords
    print("\nMatched Bullet Points from CV:")
    for point in cv_content:
        point_doc = nlp(point)
        matching_keywords = []
        for keyword in job_description_keywords:
            keyword_doc = nlp(keyword)
            similarity = point_doc.similarity(keyword_doc)
            if similarity >= 0.7:
                matching_keywords.append(keyword)
            else:
                synonyms[keyword] = list(get_synonyms(keyword))
                for synonym in synonyms[keyword]:
                    synonym_doc = nlp(synonym)
                    synonym_similarity = point_doc.similarity(synonym_doc)
                    if synonym_similarity >= 0.9:
                        matching_keywords.append(keyword)
                        break
        if matching_keywords:
            matched_points.append(point)
            similarity_scores.append(max(similarity for similarity in [point_doc.similarity(nlp(keyword)) for keyword in matching_keywords]))
    
    # Save the matched bullet points to a DOCX file
    matched_doc = docx.Document()
    for point in matched_points:
        matched_doc.add_paragraph(point)
    matched_doc.save(output_file_path)
    print("Matched bullet points saved to", output_file_path)

    # Generate and save the report to a DOCX file
    report_doc = generate_report(matched_points, job_description_keywords, similarity_scores, synonyms)
    report_doc.save(report_file_path)
    print("Report saved to", report_file_path)
